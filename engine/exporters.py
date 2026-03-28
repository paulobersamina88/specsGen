import io
import pandas as pd
from docx import Document

def to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")

def to_txt_bytes(text: str) -> bytes:
    return text.encode("utf-8")

def to_docx_bytes(project_name: str, agency_name: str, rows: list[dict]) -> bytes:
    doc = Document()
    doc.add_heading("Generated Technical Specifications", level=1)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(f"Agency / Office: {agency_name}")
    doc.add_paragraph("This document contains draft specifications for internal review and refinement.")
    for row in rows:
        doc.add_page_break()
        doc.add_heading(f"BOQ Row {row.get('row_id', '')} - {row.get('description', '')}", level=2)
        for line in str(row.get("generated_spec", "")).split("\n"):
            doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
